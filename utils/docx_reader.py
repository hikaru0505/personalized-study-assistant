"""
Extract text from a Word (.docx) file using python-docx.

DOCX files don't have a real page concept the way PDFs do (page breaks are
a rendering-time detail, not something python-docx exposes reliably), so
we treat the entire document as a single "page" for citation purposes.
This is called out explicitly in the UI/README rather than faked.
"""

from typing import List, Dict
from docx import Document


def extract_pages(docx_path: str) -> List[Dict]:
    """Returns a single-page list: [{"page": 1, "text": <full document text>}]"""
    document = Document(docx_path)

    text_parts = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)

    return [{"page": 1, "text": "\n".join(text_parts)}]


def extract_text(docx_path: str) -> str:
    """Flat full-document text, for callers that don't need page boundaries."""
    return extract_pages(docx_path)[0]["text"]
